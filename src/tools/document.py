"""
Archon Document Tools

Document reading and text extraction tools.
"""

import os
from pathlib import Path
from typing import Optional

from src.tools.base import tool, ToolResult
from src.core.types import Document, DocumentType
from src.monitoring.logger import get_logger

logger = get_logger("tools.document")


# =============================================================================
# Document Type Detection
# =============================================================================

EXTENSION_MAP = {
    ".pdf": DocumentType.PDF,
    ".docx": DocumentType.DOCX,
    ".doc": DocumentType.DOCX,
    ".txt": DocumentType.TXT,
    ".md": DocumentType.MD,
    ".markdown": DocumentType.MD,
    ".html": DocumentType.HTML,
    ".htm": DocumentType.HTML,
    ".csv": DocumentType.CSV,
    ".json": DocumentType.JSON,
}


def detect_document_type(file_path: str) -> DocumentType:
    """Detect document type from file extension."""
    ext = Path(file_path).suffix.lower()
    return EXTENSION_MAP.get(ext, DocumentType.UNKNOWN)


# =============================================================================
# Text Extraction
# =============================================================================

@tool(
    name="extract_text",
    description="Extract text content from a document file (PDF, DOCX, TXT, etc.)",
    category="document",
    parameters={
        "type": "object",
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Path to the document file",
            },
            "max_length": {
                "type": "integer",
                "description": "Maximum characters to extract (default: 50000)",
                "default": 50000,
            },
        },
        "required": ["file_path"],
    },
)
async def extract_text(file_path: str, max_length: int = 50000) -> ToolResult:
    """
    Extract text from various document formats.
    """
    if not os.path.exists(file_path):
        return ToolResult(
            success=False,
            error=f"File not found: {file_path}",
        )

    doc_type = detect_document_type(file_path)

    try:
        if doc_type == DocumentType.PDF:
            text = await _extract_pdf(file_path)
        elif doc_type == DocumentType.DOCX:
            text = await _extract_docx(file_path)
        elif doc_type in (DocumentType.TXT, DocumentType.MD):
            text = await _extract_text_file(file_path)
        elif doc_type == DocumentType.HTML:
            text = await _extract_html(file_path)
        else:
            text = await _extract_text_file(file_path)

        # Truncate if needed
        if len(text) > max_length:
            text = text[:max_length] + "\n\n[... truncated ...]"

        return ToolResult(
            success=True,
            data={
                "text": text,
                "file_path": file_path,
                "doc_type": doc_type.value,
                "char_count": len(text),
            },
        )

    except Exception as e:
        logger.error(f"Text extraction failed: {e}")
        return ToolResult(
            success=False,
            error=str(e),
        )


async def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF."""
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            text_parts.append(page.extract_text())

        return "\n\n".join(text_parts)

    except ImportError:
        # Fallback to pdfminer
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            return pdfminer_extract(file_path)
        except ImportError:
            raise ImportError("Install pypdf or pdfminer.six for PDF support")


async def _extract_docx(file_path: str) -> str:
    """Extract text from DOCX."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text_parts = []

        for para in doc.paragraphs:
            text_parts.append(para.text)

        return "\n".join(text_parts)

    except ImportError:
        raise ImportError("Install python-docx for DOCX support")


async def _extract_text_file(file_path: str) -> str:
    """Extract text from plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


async def _extract_html(file_path: str) -> str:
    """Extract text from HTML."""
    try:
        from bs4 import BeautifulSoup

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            soup = BeautifulSoup(f.read(), "html.parser")

        # Remove script and style elements
        for element in soup(["script", "style"]):
            element.decompose()

        return soup.get_text(separator="\n", strip=True)

    except ImportError:
        # Fallback to simple regex
        import re
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        # Remove tags
        clean = re.sub(r"<[^>]+>", "", content)
        return clean


# =============================================================================
# Document Reading
# =============================================================================

@tool(
    name="read_document",
    description="Read and parse a document, returning structured content.",
    category="document",
    parameters={
        "type": "object",
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Path to the document",
            },
            "include_metadata": {
                "type": "boolean",
                "description": "Include document metadata",
                "default": True,
            },
        },
        "required": ["file_path"],
    },
)
async def read_document(file_path: str, include_metadata: bool = True) -> ToolResult:
    """
    Read a document and return structured content.
    """
    if not os.path.exists(file_path):
        return ToolResult(
            success=False,
            error=f"File not found: {file_path}",
        )

    try:
        # Extract text
        text_result = await extract_text(file_path)

        if not text_result.success:
            return text_result

        text = text_result.data["text"]

        # Build document
        doc = Document(
            content=text,
            doc_type=detect_document_type(file_path),
            source=file_path,
        )

        result_data = {
            "id": doc.id,
            "content": doc.content,
            "doc_type": doc.doc_type.value,
            "source": doc.source,
            "char_count": doc.char_count,
            "word_count": doc.word_count,
        }

        if include_metadata:
            result_data["metadata"] = await _get_file_metadata(file_path)

        return ToolResult(
            success=True,
            data=result_data,
        )

    except Exception as e:
        logger.error(f"Document read failed: {e}")
        return ToolResult(
            success=False,
            error=str(e),
        )


async def _get_file_metadata(file_path: str) -> dict:
    """Get file metadata."""
    stat = os.stat(file_path)
    path = Path(file_path)

    return {
        "filename": path.name,
        "extension": path.suffix,
        "size_bytes": stat.st_size,
        "modified_time": stat.st_mtime,
        "created_time": stat.st_ctime,
    }


# =============================================================================
# URL Fetching
# =============================================================================

@tool(
    name="fetch_url",
    description="Fetch content from a URL and extract text.",
    category="document",
    parameters={
        "type": "object",
        "properties": {
            "url": {
                "type": "string",
                "description": "The URL to fetch",
            },
            "extract_main_content": {
                "type": "boolean",
                "description": "Extract main content only (removes nav, ads, etc.)",
                "default": True,
            },
        },
        "required": ["url"],
    },
)
async def fetch_url(url: str, extract_main_content: bool = True) -> ToolResult:
    """
    Fetch and parse content from a URL.
    """
    try:
        import aiohttp

        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=30) as response:
                if response.status != 200:
                    return ToolResult(
                        success=False,
                        error=f"HTTP {response.status}: {response.reason}",
                    )

                html = await response.text()

        # Parse HTML
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")

        if extract_main_content:
            # Try to find main content
            main = soup.find("main") or soup.find("article") or soup.find("body")
            if main:
                soup = main

        # Remove unwanted elements
        for element in soup(["script", "style", "nav", "footer", "header", "aside"]):
            element.decompose()

        text = soup.get_text(separator="\n", strip=True)

        # Get title
        title_tag = soup.find("title")
        title = title_tag.get_text() if title_tag else ""

        return ToolResult(
            success=True,
            data={
                "url": url,
                "title": title,
                "content": text,
                "char_count": len(text),
            },
        )

    except ImportError:
        return ToolResult(
            success=False,
            error="Install aiohttp and beautifulsoup4 for URL fetching",
        )
    except Exception as e:
        logger.error(f"URL fetch failed: {e}")
        return ToolResult(
            success=False,
            error=str(e),
        )
